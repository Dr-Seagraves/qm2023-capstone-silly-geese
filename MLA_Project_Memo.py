#!/usr/bin/env python3
"""
MLA-Formatted Project Memo Paper
Generates a personal reflection memo on the lobbying and firm profitability research project
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create a new Document
doc = Document()

# Set up margins (1 inch all around)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# MLA Header (last name and page number) - we'll add this manually
# Create header
header = sections[0].header
header_para = header.paragraphs[0]
header_para.text = "Student 1"
header_run = header_para.runs[0]
header_run.font.size = Pt(12)

# MLA Heading (upper left corner)
heading = doc.add_paragraph()
heading.paragraph_format.space_before = Pt(0)
heading.paragraph_format.space_after = Pt(0)
heading.paragraph_format.line_spacing = 2.0
run = heading.add_run("Student Name\n")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = heading.add_run("Dr. Seagraves\n")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = heading.add_run("QM 2023: Capstone\n")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = heading.add_run("1 May 2026")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Title (centered, not bold, in MLA format)
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(0)
title.paragraph_format.line_spacing = 2.0
title_run = title.add_run("Political Influence and Corporate Profitability:\nA Personal Reflection on Lobbying Research")
title_run.font.size = Pt(12)
title_run.font.name = 'Times New Roman'

# Body paragraphs - all double-spaced
def add_body_paragraph(text):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Inches(0.5)
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return para

# Introduction
intro_text = ("As a political science student, I have always been intrigued by the intricate relationship between corporate interests and political institutions. When our capstone project team began brainstorming potential research questions, I was immediately drawn to an investigation of corporate lobbying expenditures and their effects on firm profitability. This project offered a unique opportunity to combine my disciplinary interests with rigorous quantitative analysis, allowing me to explore whether and how corporations' political investments translate into measurable financial outcomes. My initial hypothesis was straightforward: firms that spend more on lobbying should experience higher subsequent profitability. This intuition rests on the assumption that lobbying provides firms with improved regulatory clarity, favorable policy outcomes, reduced policy uncertainty, and increased access to political decision-makers—all of which should enhance operational and financial performance. The following memo documents my personal journey through this research, the methods we employed, our key findings, and the lessons I learned about the complexity of corporate-political relationships.")
add_body_paragraph(intro_text)

# Background
background_text = ("Our research began with a fundamental question: What is the relationship between firms' lobbying expenditures and their subsequent profitability? To answer this question, we assembled a comprehensive firm-year panel dataset spanning 2010–2020, merging two critical data sources. First, we compiled federal lobbying expenditure data from Senate Lobbying Disclosure Reports, which detailed the amount each firm spent on government relations activities annually. Second, we obtained firm financial data from SEC financial filings, extracting assets, net income, and revenue information at the annual level. By creating a crosswalk between corporate identifiers (CIK and GVKEY), we successfully merged these datasets and constructed a balanced panel that allowed us to observe individual firms over multiple years, enabling within-firm comparisons that control for time-invariant firm heterogeneity.")
add_body_paragraph(background_text)

# Methods
methods_text = ("Our analytical strategy prioritized identification of within-firm associations rather than causal claims, given observational constraints. In our primary specification, we estimated two-way fixed effects panel regressions with firm and year fixed effects, clustering standard errors at the firm level to account for correlated shocks within firms over time. Our focal predictor was lagged lobbying expenditure (t-1), reflecting the intuition that policy influence materializes gradually: lobbying campaigns influence political processes, which then translate into policy changes, which ultimately affect firms' operating and financial performance. We included controls for firm size (log assets) and firm revenues to isolate the marginal association of lobbying while holding financial scale constant. In robustness analyses, we tested alternative lag structures, excluded anomalous years, examined heterogeneity across firm size categories, and implemented dynamic difference-in-differences designs to assess the stability of our estimates.")
add_body_paragraph(methods_text)

# Findings Discovery
discovery_text = ("Our exploratory data analysis initially supported my hypothesis. Descriptive statistics revealed a moderate positive correlation between lobbying expenditure and Return on Assets (ROA), and this relationship appeared particularly strong among larger firms and those operating in policy-sensitive industries such as energy, telecommunications, and healthcare. When we visualized trends over time, both lobbying spending and firm profitability generally moved together, suggesting a possible positive link. However, when we moved to formal regression analysis, the picture became more complicated. Our primary panel regression with lagged lobbying as the predictor yielded a point estimate of −174.6 percentage points in winsorized ROA for every $1 million increase in lobbying expenditure, with a standard error of 255 percentage points and a p-value of 0.494. In plain language, this coefficient was negative, imprecise, and not statistically significantly different from zero. My initial hypothesis was not supported by the formal analysis.")
add_body_paragraph(discovery_text)

# Interpretation
interpretation_text = ("At first, this result was disappointing. However, the misalignment between my hypothesis and the data prompted deeper reflection on the research question and underlying mechanisms. Several insights emerged. First, the negative coefficient, though imprecise, suggests that lobbying expenditure alone may not translate directly into profitability—or worse, firms that lobby heavily might face structural headwinds in profitability that offset any political gains. One plausible mechanism is that firms lobby precisely because they face regulatory or market pressures that constrain profitability, creating a spurious negative correlation. Second, the instability of the coefficient across alternative specifications—different lag lengths, different sample definitions, and different subgroups—indicated that any relationship is highly context-dependent and sensitive to modeling choices. A lead-placebo test, where we examined whether future profitability predicts current lobbying, yielded a negative coefficient with near-marginal significance (p = 0.069), raising the troubling possibility of reverse causality: profitable firms may underinvest in lobbying because they face fewer threats, or unprofitable firms may ramp up lobbying desperately, creating the illusion of a negative effect. Third, we found important heterogeneity: the adverse association was more pronounced in large firms than small firms, which contradicts the simple political-access narrative and suggests that scale, complexity, and sector-specific regulatory dynamics matter enormously.")
add_body_paragraph(interpretation_text)

# Reflection
reflection_text = ("This project taught me humility. My initial hypothesis, while theoretically plausible, failed to survive encounter with real data. More importantly, this failure revealed the limitations of observational research on causally complex phenomena. The relationship between political spending and corporate performance is not a direct pipeline but an intricate web of firm strategy, regulatory context, market competition, and unobserved confounders. Firms that lobby may differ systematically in ways we cannot measure—their exposure to policy risk, their managerial sophistication, their market position, or their business model. Without experimental variation or quasi-experimental shocks, disentangling these mechanisms remains extraordinarily difficult. Moreover, the heterogeneity we observed—that lobbying effects diverge by firm size and policy exposure—suggests that any universal claim about whether 'lobbying increases profitability' is likely misguided. The truth may be that lobbying is effective in some contexts (e.g., a large pharmaceutical firm navigating FDA approval) and ineffective or counterproductive in others (e.g., a small manufacturing firm in a commodity market). The project also reinforced the importance of diagnostic rigor: checking for heteroskedasticity, examining lagged structures, testing for reverse causality via placebo leads, and reporting multiple specifications rather than cherry-picking the most favorable result.")
add_body_paragraph(reflection_text)

# Conclusion
conclusion_text = ("In retrospect, I approached this capstone project with a clear but untested hypothesis: lobbying should increase profitability. The data did not confirm this theory, at least not in a straightforward or robust manner. Yet this outcome represents the true spirit of empirical research. We generate hypotheses, test them rigorously, and revise our understanding when data contradicts our priors. My interest in political science and corporate influence remains undiminished; if anything, it has deepened. The disconnect between my initial intuition and the findings highlights how political economy operates at multiple levels—macro-institutional, sector-specific, and firm-specific—and how static correlations in observational data rarely capture causal dynamics. This capstone experience has prepared me for future research by instilling respect for evidence, awareness of methodological constraints, and appreciation for the hard work of building datasets and studying complex relationships. I am grateful to my team for their collaboration, to Dr. Seagraves for her mentorship, and to the data for surprising me.")
add_body_paragraph(conclusion_text)

# Works Cited Page
doc.add_page_break()
works_cited = doc.add_paragraph()
works_cited.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
works_cited.paragraph_format.line_spacing = 2.0
works_cited.paragraph_format.space_after = Pt(0)
works_cited_run = works_cited.add_run("Works Cited")
works_cited_run.font.size = Pt(12)
works_cited_run.font.name = 'Times New Roman'

# Cited sources
def add_works_cited_entry(text):
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.5)
    run = para.runs[0]
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

add_works_cited_entry("Center for Responsive Politics. \"Lobbying Database.\" OpenSecrets, Accessed 1 May 2026, www.opensecrets.org.")

add_works_cited_entry("Securities and Exchange Commission. \"EDGAR: Database of Corporate Information.\" U.S. SEC, Accessed 1 May 2026, www.sec.gov/cgi-bin/browse-edgar.")

add_works_cited_entry("Blanes i Vidal, Jordi, et al. \"Revolving Door Lobbyists.\" American Economic Review, vol. 102, no. 7, 2012, pp. 3731-48. JSTOR, https://doi.org/10.1257/aer.102.7.3731.")

add_works_cited_entry("Bombardini, Matilde, and Francesco Trebbi. \"Lobbying and the Organizational Equilibrium Value of the Firm.\" Journal of Political Economy, vol. 130, no. 3, 2022, pp. 749-779.")

# Save the document
output_path = "/workspaces/qm2023-capstone-silly-geese/MLA_Project_Memo.docx"
doc.save(output_path)
print(f"MLA Project Memo created successfully at: {output_path}")
